"""
Gerador de Relatórios Técnicos de Medição - AIR MINAS AR CONDICIONADO LTDA
Monta um .docx profissional no padrão exigido pelo DER-ES a partir dos dados
informados pelo usuário (contrato, medição, itens medidos, segurança do
trabalho, curva S e responsável técnico).
"""
from __future__ import annotations

import io
import os
import uuid
from datetime import datetime

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Paleta de cores extraída dos relatórios de referência da AIR MINAS
# ---------------------------------------------------------------------------
DARK_BLUE = "1F6EB1"      # cabeçalho das tabelas (fundo)
TEXT_BLUE = "0E4D79"      # texto sobre fundo claro
LIGHT_1 = "E9F4FB"
LIGHT_2 = "D8EBF6"
BORDER_BLUE = "5B9AD4"
WHITE = "FFFFFF"
RED_EXEC = "C0504D"

FONT_NAME = "Arial"

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_LOGO = os.path.join(BASE_DIR, "static", "img", "air_minas_logo.png")

# Ordem exigida pelo XSD do OOXML para os filhos de <w:tcPr> e <w:pPr>.
# Elementos precisam ser inseridos respeitando essa sequência, senão o
# documento fica inválido mesmo que o Word normalmente tolere a abertura.
_TCPR_ORDER = [qn(t) for t in (
    "w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge", "w:tcBorders",
    "w:shd", "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText",
    "w:vAlign", "w:hideMark",
)]
_PPR_ORDER = [qn(t) for t in (
    "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore", "w:framePr",
    "w:widowControl", "w:numPr", "w:suppressLineNumbers", "w:pBdr", "w:shd",
    "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
    "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
)]


def _insert_ordered(parent, new_el, order):
    """Insere new_el em parent respeitando a sequência do XSD em order."""
    new_idx = order.index(new_el.tag)
    for child in parent:
        child_tag = child.tag
        if child_tag in order and order.index(child_tag) > new_idx:
            child.addprevious(new_el)
            return
    parent.append(new_el)


# ---------------------------------------------------------------------------
# Helpers de baixo nível (OOXML) para reproduzir a formatação das planilhas
# ---------------------------------------------------------------------------
def _set_cell_background(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    _insert_ordered(tcPr, shd, _TCPR_ORDER)


def _set_cell_borders(cell, color=BORDER_BLUE, size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    _insert_ordered(tcPr, borders, _TCPR_ORDER)


def _set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
    for idx, width in enumerate(widths_cm):
        table.columns[idx].width = Cm(width)


def _cell_text(cell, text, bold=False, color=TEXT_BLUE, size=9, align="center",
               font=FONT_NAME, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.text = ""
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text) if text not in (None, "") else "-")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def _add_heading(doc, text, size=13, color=DARK_BLUE, space_before=14, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor.from_string(color)
    # linha inferior fina como separador
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BORDER_BLUE)
    pBdr.append(bottom)
    _insert_ordered(pPr, pBdr, _PPR_ORDER)
    return p


def _fmt_money(value) -> str:
    try:
        v = float(value)
    except (TypeError, ValueError):
        return "-"
    s = f"{v:,.2f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {s}"


def _fmt_num(value, casas=2) -> str:
    try:
        v = float(value)
    except (TypeError, ValueError):
        return "-"
    s = f"{v:,.{casas}f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return s


def _fmt_date(value) -> str:
    if not value:
        return "-"
    if isinstance(value, str):
        for fmt_in in ("%Y-%m-%d", "%d/%m/%Y"):
            try:
                dt = datetime.strptime(value, fmt_in)
                return dt.strftime("%d/%m/%Y")
            except ValueError:
                continue
        return value
    return value.strftime("%d/%m/%Y")


# ---------------------------------------------------------------------------
# Cabeçalho / rodapé
# ---------------------------------------------------------------------------
def _build_header_footer(section, data, logo_path):
    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        p.text = ""
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if logo_path and os.path.exists(logo_path):
        run = hp.add_run()
        run.add_picture(logo_path, width=Cm(4.2))

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p.text = ""
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = fp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), BORDER_BLUE)
    pBdr.append(top)
    _insert_ordered(pPr, pBdr, _PPR_ORDER)

    contratada = data.get("contrato", {}).get("contratada", "AIR MINAS AR CONDICIONADO LTDA")
    contrato_num = data.get("contrato", {}).get("contrato", "")
    medicao_num = data.get("medicao", {}).get("numero", "")
    periodo = data.get("medicao", {}).get("periodo_faturamento", "")
    texto = f"{contratada} · Contrato {contrato_num} · RELATÓRIO TÉCNICO – {medicao_num}ª Medição · {periodo}"
    run = fp.add_run(texto)
    run.font.size = Pt(7.5)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor.from_string(TEXT_BLUE)

    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = fp2.add_run("Página ")
    run2.font.size = Pt(7.5)
    run2.font.name = FONT_NAME
    run2.font.color.rgb = RGBColor.from_string(TEXT_BLUE)
    _add_field(fp2, "PAGE")
    run3 = fp2.add_run(" de ")
    run3.font.size = Pt(7.5)
    run3.font.name = FONT_NAME
    run3.font.color.rgb = RGBColor.from_string(TEXT_BLUE)
    _add_field(fp2, "NUMPAGES")


def _add_field(paragraph, field_code):
    run = paragraph.add_run()
    run.font.size = Pt(7.5)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor.from_string(TEXT_BLUE)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


# ---------------------------------------------------------------------------
# Tabela chave/valor (padrão dos cabeçalhos de contrato / medição)
# ---------------------------------------------------------------------------
def _info_table(doc, pairs, col_widths=(4.5, 9.0)):
    table = doc.add_table(rows=len(pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(pairs):
        light = LIGHT_1 if i % 2 == 0 else LIGHT_2
        label_cell, value_cell = table.rows[i].cells
        _set_cell_background(label_cell, light)
        _set_cell_background(value_cell, light)
        _set_cell_borders(label_cell)
        _set_cell_borders(value_cell)
        _cell_text(label_cell, label, bold=True, color=TEXT_BLUE, align="center")
        _cell_text(value_cell, value, bold=False, color=TEXT_BLUE, align="center")
    _set_col_widths(table, col_widths)
    return table


def _title_row_table(doc, header_left, header_right, col_widths=(4.5, 9.0)):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    a, b = table.rows[0].cells
    for c in (a, b):
        _set_cell_background(c, DARK_BLUE)
        _set_cell_borders(c)
    _cell_text(a, header_left, bold=True, color=WHITE, align="center")
    _cell_text(b, header_right, bold=True, color=WHITE, align="center")
    _set_col_widths(table, col_widths)
    return table


# ---------------------------------------------------------------------------
# Tabela de itens medidos
# ---------------------------------------------------------------------------
def _items_table(doc, itens):
    headers = [
        "Item", "Descrição", "Und.", "Qtd.\nContratada", "Qtd.\nAcum. Anterior",
        "Qtd.\nMedição Atual", "Qtd.\nAcum. Atual", "Preço Unit. (R$)", "Valor Medido (R$)",
    ]
    widths = [1.4, 5.2, 1.2, 1.9, 1.9, 1.9, 1.9, 2.2, 2.4]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(headers):
        cell = table.rows[0].cells[j]
        _set_cell_background(cell, DARK_BLUE)
        _set_cell_borders(cell)
        _cell_text(cell, htxt, bold=True, color=WHITE, size=8.5, align="center")

    total_valor = 0.0
    for i, item in enumerate(itens):
        qtd_contratada = float(item.get("qtd_contratada") or 0)
        qtd_anterior = float(item.get("qtd_anterior") or 0)
        qtd_atual = float(item.get("qtd_atual") or 0)
        preco_unit = float(item.get("preco_unitario") or 0)
        qtd_acumulada = qtd_anterior + qtd_atual
        valor_medido = qtd_atual * preco_unit
        total_valor += valor_medido

        row_cells = table.add_row().cells
        light = LIGHT_1 if i % 2 == 0 else LIGHT_2
        values = [
            item.get("item", ""),
            item.get("descricao", ""),
            item.get("unidade", ""),
            _fmt_num(qtd_contratada),
            _fmt_num(qtd_anterior),
            _fmt_num(qtd_atual),
            _fmt_num(qtd_acumulada),
            _fmt_num(preco_unit),
            _fmt_num(valor_medido),
        ]
        for j, val in enumerate(values):
            cell = row_cells[j]
            _set_cell_background(cell, light)
            _set_cell_borders(cell)
            align = "left" if j == 1 else "center"
            _cell_text(cell, val, bold=(j == 0), color=TEXT_BLUE, size=8.5, align=align)

    # linha de total
    total_cells = table.add_row().cells
    for c in total_cells:
        _set_cell_background(c, DARK_BLUE)
        _set_cell_borders(c)
    _cell_text(total_cells[0], "TOTAL", bold=True, color=WHITE, size=9, align="left")
    for j in range(1, 8):
        _cell_text(total_cells[j], "", color=WHITE)
    _cell_text(total_cells[8], _fmt_num(total_valor), bold=True, color=WHITE, size=9)

    _set_col_widths(table, widths)
    return table, total_valor


def _seguranca_table(doc, seg):
    pairs = [
        ("Acidentes com afastamento", seg.get("com_afastamento", 0)),
        ("Acidentes sem afastamento", seg.get("sem_afastamento", 0)),
        ("Acidente típico", seg.get("tipo_tipico", 0)),
        ("Acidente de trajeto", seg.get("tipo_trajeto", 0)),
        ("Doença profissional", seg.get("tipo_doenca", 0)),
        ("HHT (Homens-Hora Trabalhadas)", seg.get("hht", 0)),
        ("Taxa de gravidade", seg.get("taxa_gravidade", 0)),
    ]
    _info_table(doc, pairs, col_widths=(7.0, 6.5))


def _curva_s_chart(curva_s):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    meses = [c.get("mes", "") for c in curva_s]
    previsto = [float(c.get("previsto_acumulado") or 0) for c in curva_s]
    executado = [float(c.get("executado_acumulado") or 0) for c in curva_s]

    fig, ax = plt.subplots(figsize=(9.2, 3.6), dpi=200)
    ax.plot(meses, previsto, marker="o", color="#1F6EB1", linewidth=2, label="Previsto Acumulado")
    ax.plot(meses, executado, marker="o", color="#C0504D", linewidth=2, label="Executado Acumulado")
    ax.set_title("Curva S – Previsto x Executado", fontsize=11, color="#0E4D79")
    ax.set_ylabel("R$ Acumulado", fontsize=8)
    ax.tick_params(axis="x", labelrotation=45, labelsize=7)
    ax.tick_params(axis="y", labelsize=7)
    ax.grid(True, linestyle="--", linewidth=0.4, alpha=0.6)
    ax.legend(fontsize=8, loc="upper left")
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Função principal
# ---------------------------------------------------------------------------
def build_report(data: dict, logo_bytes: bytes | None = None,
                  assinatura_bytes: bytes | None = None) -> io.BytesIO:
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    logo_path = DEFAULT_LOGO
    tmp_logo = None
    if logo_bytes:
        tmp_logo = os.path.join(BASE_DIR, "app_data", f"_tmp_logo_{uuid.uuid4().hex}.png")
        with open(tmp_logo, "wb") as f:
            f.write(logo_bytes)
        logo_path = tmp_logo

    _build_header_footer(section, data, logo_path)

    contrato = data.get("contrato", {})
    medicao = data.get("medicao", {})

    # Página de rosto -------------------------------------------------
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run(data.get("titulo") or "RELATÓRIO TÉCNICO")
    trun.font.name = FONT_NAME
    trun.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    trun.font.size = Pt(20)
    trun.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run(f"{medicao.get('numero','')}ª Medição – {medicao.get('periodo_faturamento','')}")
    srun.font.size = Pt(12)
    srun.font.name = FONT_NAME
    srun.italic = True
    srun.font.color.rgb = RGBColor.from_string(TEXT_BLUE)
    doc.add_paragraph()

    obj_pairs = [
        ("OBJETO", contrato.get("objeto", "")),
        ("CONTRATADA", contrato.get("contratada", "")),
        ("CONTRATO", contrato.get("contrato", "")),
        ("ORDEM DE SERVIÇO", contrato.get("os", "")),
        ("VALOR CONTRATUAL", _fmt_money(contrato.get("valor_contratual"))),
        ("INÍCIO DOS SERVIÇOS", _fmt_date(contrato.get("inicio"))),
        ("TÉRMINO PREVISTO", _fmt_date(contrato.get("termino"))),
        ("PRAZO DE EXECUÇÃO", contrato.get("prazo", "")),
    ]
    _info_table(doc, obj_pairs)
    doc.add_paragraph()

    med_pairs = [
        ("MEDIÇÃO", f"{medicao.get('numero','')}ª Medição"),
        ("VALOR DA MEDIÇÃO", _fmt_money(medicao.get("valor_medicao"))),
        ("PERÍODO (FATURAMENTO)", medicao.get("periodo_faturamento", "")),
        ("PERÍODO (DOCUMENTAÇÃO)", medicao.get("periodo_documentacao", "")),
    ]
    _info_table(doc, med_pairs)
    doc.add_paragraph()

    fisc = doc.add_paragraph()
    fisc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fisc.add_run(f"Fiscalização: {contrato.get('fiscalizacao', 'DER-ES')}")
    frun.bold = True
    frun.font.size = Pt(10.5)
    frun.font.name = FONT_NAME
    frun.font.color.rgb = RGBColor.from_string(TEXT_BLUE)

    ident = doc.add_paragraph()
    ident.alignment = WD_ALIGN_PARAGRAPH.CENTER
    irun = ident.add_run(
        f"{contrato.get('contratada','')} · Contrato {contrato.get('contrato','')} · "
        f"Vitória/ES · {medicao.get('periodo_faturamento','')}"
    )
    irun.font.size = Pt(9)
    irun.font.name = FONT_NAME
    irun.font.color.rgb = RGBColor.from_string(TEXT_BLUE)

    doc.add_page_break()

    # Itens medidos -----------------------------------------------------
    _add_heading(doc, "1. MEDIÇÃO DOS SERVIÇOS EXECUTADOS")
    itens = data.get("itens", [])
    if itens:
        table, total_valor = _items_table(doc, itens)
    else:
        p = doc.add_paragraph("Nenhum item informado para esta medição.")
        p.italic = True
        total_valor = 0.0

    resumo = doc.add_paragraph()
    resumo.paragraph_format.space_before = Pt(10)
    rrun = resumo.add_run(f"Valor total medido nesta medição: {_fmt_money(total_valor)}")
    rrun.bold = True
    rrun.font.name = FONT_NAME
    rrun.font.size = Pt(10.5)
    rrun.font.color.rgb = RGBColor.from_string(TEXT_BLUE)

    # Curva S -------------------------------------------------------------
    curva_s = data.get("curva_s") or []
    if curva_s:
        _add_heading(doc, "2. CRONOGRAMA FÍSICO-FINANCEIRO – CURVA S")
        chart_buf = _curva_s_chart(curva_s)
        doc.add_picture(chart_buf, width=Cm(16.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Segurança do trabalho -------------------------------------------
    seguranca = data.get("seguranca")
    if seguranca:
        _add_heading(doc, f"{3 if curva_s else 2}. SEGURANÇA DO TRABALHO")
        _seguranca_table(doc, seguranca)

    # Observações --------------------------------------------------------
    observacoes = data.get("observacoes")
    if observacoes:
        n = 2 + (1 if curva_s else 0) + (1 if seguranca else 0)
        _add_heading(doc, f"{n}. OBSERVAÇÕES E CONSIDERAÇÕES FINAIS")
        for paragrafo in str(observacoes).split("\n"):
            if paragrafo.strip():
                p = doc.add_paragraph(paragrafo.strip())
                p.paragraph_format.space_after = Pt(6)
                for r in p.runs:
                    r.font.name = FONT_NAME
                    r.font.size = Pt(10.5)

    # Assinatura -----------------------------------------------------------
    doc.add_paragraph()
    doc.add_paragraph()
    responsavel = data.get("responsavel", {})
    resp_p = doc.add_paragraph()
    resp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp_run = resp_p.add_run("Responsável pela elaboração:")
    rp_run.font.name = FONT_NAME
    rp_run.font.size = Pt(10)
    rp_run.font.color.rgb = RGBColor.from_string(TEXT_BLUE)

    tmp_sig = None
    if assinatura_bytes:
        tmp_sig = os.path.join(BASE_DIR, "app_data", f"_tmp_assinatura_{uuid.uuid4().hex}.png")
        with open(tmp_sig, "wb") as f:
            f.write(assinatura_bytes)
        sig_p = doc.add_paragraph()
        sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_p.add_run().add_picture(tmp_sig, width=Cm(5.5))
    else:
        doc.add_paragraph()
        line_p = doc.add_paragraph("_" * 40)
        line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    nome_p = doc.add_paragraph()
    nome_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nrun = nome_p.add_run(responsavel.get("nome", ""))
    nrun.bold = True
    nrun.font.name = FONT_NAME
    nrun.font.size = Pt(10.5)
    nrun.font.color.rgb = RGBColor.from_string(TEXT_BLUE)

    cargo_p = doc.add_paragraph()
    cargo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cargo_p.add_run(
        " · ".join(x for x in [responsavel.get("cargo", ""), responsavel.get("crea", "")] if x)
    )
    crun.font.name = FONT_NAME
    crun.font.size = Pt(9.5)
    crun.font.color.rgb = RGBColor.from_string(TEXT_BLUE)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    for tmp_file in (tmp_logo, tmp_sig):
        if tmp_file and os.path.exists(tmp_file):
            os.remove(tmp_file)

    return output
